from docx import Document
import json

doc = Document("d:/Downloads/IMPACT Product details (1).docx")
output = []

for element in doc.element.body:
    if element.tag.endswith('p'): # Paragraph
        text = "".join(node.text for node in element.iter() if node.tag.endswith('t') and node.text)
        if text.strip():
            output.append({"type": "paragraph", "text": text.strip()})
    elif element.tag.endswith('tbl'): # Table
        # We parse the table using python-docx wrapper instead of raw xml for ease
        pass

# A cleaner way is to iterate over block level elements using a helper function, 
# but python-docx has `doc.paragraphs` and `doc.tables`.
# The problem is that doc.paragraphs only gives paragraphs outside tables. 
# They are not in sequential order of the document structure (interleaved).
# We can use the parent element to get them in order.

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    """
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

items = []
for block in iter_block_items(doc):
    if block.__class__.__name__ == 'Paragraph':
        text = block.text.strip()
        if text:
            items.append({"type": "paragraph", "text": text})
    elif block.__class__.__name__ == 'Table':
        table_data = []
        for row in block.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        items.append({"type": "table", "data": table_data})

with open("d:/New-Impact-Care/docx_extract/structure.json", "w", encoding="utf-8") as f:
    json.dump(items, f, indent=2, ensure_ascii=False)

print(f"Extracted {len(items)} items to structure.json")
